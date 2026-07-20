from pathlib import Path
import zipfile

from docx import Document


src = Path(r"C:\Users\navni\Downloads\TicketRouter_Submission.docx")
out = Path(r"C:\Users\navni\OneDrive\Documents\Week 5 Project\tmp\linkedin_carousel")
media = out / "media"
out.mkdir(parents=True, exist_ok=True)
media.mkdir(parents=True, exist_ok=True)

doc = Document(src)
lines = []
for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    if text:
        lines.append(text)

for table in doc.tables:
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " | ") for cell in row.cells]
        if any(cells):
            lines.append("TABLE: " + " | ".join(cells))

(out / "extracted_text.txt").write_text("\n".join(lines), encoding="utf-8")

with zipfile.ZipFile(src) as archive:
    for name in archive.namelist():
        if name.startswith("word/media/"):
            target = media / Path(name).name
            target.write_bytes(archive.read(name))

print("lines", len(lines))
print("media", len(list(media.iterdir())))
print(out / "extracted_text.txt")
for item in sorted(media.iterdir()):
    print(item.name, item.stat().st_size)
